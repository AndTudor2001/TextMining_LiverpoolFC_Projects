import sqlite3
from docx import Document
import re
from collections import Counter
import matplotlib.pyplot as plt

# Funcție pentru a extrage primul paragraf dintr-un text
def extract_first_paragraph(text):
    paragraphs = text.split('\n')  # Împărțim textul pe paragrafe
    return paragraphs[0].strip() if paragraphs else text

# Funcție pentru a extrage prima propoziție dintr-un text
def extract_first_sentence(text):
    sentences = text.split('. ')  # Împărțim textul pe propoziții
    return sentences[0].strip() if sentences else text

# Funcție pentru a exporta datele combinate din tabele într-un document Word
def export_relevance_and_sentiment_to_word():
    # Creăm un document nou
    doc = Document()
    doc.add_heading('Relevanța și Sentimentul articolelor', 0)

    # Conectare la baza de date
    conn = sqlite3.connect('articoleLiverpool.db')
    cursor = conn.cursor()

    # Obținem datele din ambele tabele și le unim pe baza article_id
    cursor.execute('''
        SELECT relevanta_articole.article_id, relevanta_articole.relevance_score, sentimente_articole.sentiment
        FROM relevanta_articole
        LEFT JOIN sentimente_articole ON relevanta_articole.article_id = sentimente_articole.article_id
    ''')
    combined_results = cursor.fetchall()

    # Creăm un tabel în documentul Word pentru a afișa datele combinate
    doc.add_heading('Tabel Relevanță și Sentiment Articole', level=1)
    table_combined = doc.add_table(rows=1, cols=3)  # Trei coloane: ID, Relevance, Sentiment
    hdr_cells = table_combined.rows[0].cells
    hdr_cells[0].text = 'Article ID'
    hdr_cells[1].text = 'First Paragraph of Relevance Rating'
    hdr_cells[2].text = 'First Sentence of Sentiment'

    # Adăugăm datele în tabel
    for row in combined_results:
        row_cells = table_combined.add_row().cells
        row_cells[0].text = str(row[0])  # Article ID

        # Extragem primul paragraf din Relevance Score (dacă există)
        if row[1]:
            first_paragraph = extract_first_paragraph(row[1])
        else:
            first_paragraph = 'N/A'
        row_cells[1].text = first_paragraph

        # Extragem prima propoziție din Sentiment (dacă există)
        if row[2]:
            first_sentence = extract_first_sentence(row[2])
        else:
            first_sentence = 'N/A'
        row_cells[2].text = first_sentence

    # Salvăm documentul Word
    doc.save('Relevanta_Sentiment_Articole.docx')

    # Închidem conexiunea la baza de date
    conn.close()

# Funcție pentru a crea un grafic al celor mai frecvente structuri controversate
def plot_controversial_topics():
    # Conectare la baza de date
    conn = sqlite3.connect('articoleLiverpool.db')
    cursor = conn.cursor()

    # Selectăm datele din tabelul 'controversial_content'
    cursor.execute("SELECT analysis FROM controversial_content")
    analysis_results = cursor.fetchall()

    # Funcție pentru a extrage structurile dintre două stelute
    def extract_starred_structures(text):
        # Folosim regex pentru a găsi toate structurile dintre două stelute (**)
        return re.findall(r"\*\*(.*?)\*\*", text)

    # Contorizăm frecvențele structurilor extrase
    structure_counter = Counter()

    for row in analysis_results:
        analysis_text = row[0]
        structures = extract_starred_structures(analysis_text)  # Extragem structurile dintre două stelute
        structure_counter.update(structures)  # Actualizăm contorul de frecvență

    # Obținem Top 10 cele mai frecvente structuri
    top_10_structures = structure_counter.most_common(10)

    # Separăm structurile și frecvențele pentru a le folosi în grafic
    structures = [item[0] for item in top_10_structures]
    frequencies = [item[1] for item in top_10_structures]

    # Creăm un grafic cu bare
    plt.figure(figsize=(10, 6))
    plt.barh(structures, frequencies)
    plt.xlabel('Frequency')
    plt.ylabel('Controversial Topics')
    plt.title('Top 10 Controversial Topics (Starred Structures)')
    plt.gca().invert_yaxis()  # Invertim axa Y pentru a afișa cele mai frecvente structuri sus
    plt.tight_layout()
    plt.show()


    conn.close()

# Apelăm funcțiile pentru a genera documentul Word și graficul
export_relevance_and_sentiment_to_word()
plot_controversial_topics()

# Funcție pentru a exporta datele din tabelul 'controversial_content' într-un document Word
def export_to_word():
    conn = sqlite3.connect('articoleLiverpool.db')
    cursor = conn.cursor()

    cursor.execute("SELECT article_id, analysis FROM controversial_content")
    analysis_results = cursor.fetchall()

    doc = Document()
    doc.add_heading('Controversial Content Analysis', 0)

    for row in analysis_results:
        article_id = row[0]
        analysis = row[1]
        doc.add_paragraph(f"Article ID: {article_id}")
        doc.add_paragraph(f"Analysis: {analysis}")
        doc.add_paragraph('')  # Linie goală între articole

    doc.save('Controversial_Content_Analysis.docx')
    conn.close()


export_to_word()


conn = sqlite3.connect('articoleLiverpool.db')
cursor = conn.cursor()

# Selectăm scorurile de relevanță din tabelul 'relevanta_articole'
cursor.execute("SELECT article_id, relevance_score FROM relevanta_articole")
relevance_data = cursor.fetchall()


conn.close()

# Filtrăm primul număr găsit în relevanța articolului și contorizăm scorurile
relevance_scores = []

for item in relevance_data:
    article_id, relevance_score_text = item
    # Extragem primul număr din text folosind expresii regulate
    match = re.search(r'\d+(\.\d+)?', relevance_score_text)
    if match:
        # Convertim rezultatul găsit la float și adăugăm la lista de scoruri
        relevance_score = float(match.group(0))
        relevance_scores.append(int(relevance_score))  # Convertim la întreg pentru frecvență

# Contorizăm frecvențele fiecărui scor
score_counts = Counter(relevance_scores)

# Sortăm scorurile pentru a le afișa într-o ordine crescătoare în grafic
sorted_scores = sorted(score_counts.keys())
frequencies = [score_counts[score] for score in sorted_scores]

# Creăm graficul cu bare pentru frecvențele scorurilor de relevanță
plt.figure(figsize=(10, 6))
plt.bar(sorted_scores, frequencies)
plt.xlabel('Relevance Score')
plt.ylabel('Number of Articles')
plt.title('Frequency of Relevance Scores for Articles on Liverpool FC')
plt.xticks(sorted_scores)
plt.tight_layout()
plt.show()
